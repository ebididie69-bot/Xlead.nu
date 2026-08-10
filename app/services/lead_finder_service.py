#!/usr/bin/env python3
"""
lead_finder.py — Find local businesses WITHOUT a website, by niche + region/country.
Ranks results by review count/rating, and pulls WhatsApp number, email, and
Telegram contact where available. Runs on Termux (Android) or any Python 3
environment. Single file.

NEW: After finding leads, optionally push them directly to your LeadForge AI
dashboard so they appear immediately in the web app, ready for website
generation and outreach — no manual copy-paste or CSV import needed.

WHAT IT DOES
------------
1. You pick one or more niches from the built-in list (80+ niches), or type
   any free-form niche and it falls back to a generic OSM category search.
2. You pick continent(s), then country/countries, then optionally narrow to
   specific cities (or search the whole country).
3. It queries OpenStreetMap (Overpass API, free, no key) for matching
   businesses. If OSM finds few results AND you have a Google Places key,
   it automatically falls back to Google for that location.
4. It filters out businesses that already have a working website.
5. Writes a Word (.docx) report ranked by review count + rating.
6. Optionally pushes results straight to your LeadForge AI web app.

SETUP (Termux)
---------------
    pkg install python
    pip install requests python-docx --break-system-packages

OPTIONAL (for Google Places fallback + LeadForge push):
    Edit the CONFIG section below with your keys.

RUN
---
    python lead_finder.py
"""

import os
import re
import sys
import time
import unicodedata
from concurrent.futures import ThreadPoolExecutor, as_completed
from datetime import datetime

import requests

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Missing dependency. Run:\n  pip install python-docx --break-system-packages")
    sys.exit(1)

# ===========================================================================
# CONFIG — edit these to enable optional features
# ===========================================================================

# Your Google Places API key (optional).
# Get one at https://console.cloud.google.com/ with "Places API" enabled.
# Leave blank to use OSM only (free, but no ratings/reviews).
GOOGLE_API_KEY = ""

# Your LeadForge AI backend URL and import token.
# Set LEADFORGE_URL to your Render backend URL.
# Set LEADFORGE_TOKEN to the value of IMPORT_API_KEY you added on Render.
# Leave both blank to skip the push step entirely.
LEADFORGE_URL = "https://xlead-nu-3.onrender.com"
LEADFORGE_TOKEN = "ebified"

# Auto-push without asking each time (True = always push if configured,
# False = ask after every search).
AUTO_PUSH = False

# ===========================================================================

NOMINATIM_URL = "https://nominatim.openstreetmap.org/search"
OVERPASS_URLS = [
    "https://overpass-api.de/api/interpreter",
    "https://overpass.kumi.systems/api/interpreter",
    "https://lz4.overpass-api.de/api/interpreter",
]
GOOGLE_TEXTSEARCH_URL = "https://maps.googleapis.com/maps/api/place/textsearch/json"
GOOGLE_DETAILS_URL = "https://maps.googleapis.com/maps/api/place/details/json"
HEADERS = {"User-Agent": "lead-finder-script/1.0 (personal use)"}

MIN_OSM_RESULTS_BEFORE_FALLBACK = 3
GOOGLE_MAX_RESULTS_PER_LOCATION = 20
GOOGLE_DETAIL_WORKERS = 5

SESSION = requests.Session()
SESSION.headers.update(HEADERS)

# ---------------------------------------------------------------------------
# Niche map (80+ niches -> OSM tags)
# ---------------------------------------------------------------------------
NICHE_MAP = {
    "restaurant": [("amenity", "restaurant")],
    "cafe": [("amenity", "cafe")],
    "coffee shop": [("amenity", "cafe")],
    "bakery": [("shop", "bakery")],
    "bar": [("amenity", "bar")],
    "pub": [("amenity", "pub")],
    "car rental": [("amenity", "car_rental"), ("shop", "car_rental")],
    "car repair": [("shop", "car_repair")],
    "hotel": [("tourism", "hotel")],
    "guesthouse": [("tourism", "guest_house")],
    "gym": [("leisure", "fitness_centre")],
    "hairdresser": [("shop", "hairdresser")],
    "barber": [("shop", "hairdresser")],
    "beauty salon": [("shop", "beauty")],
    "dentist": [("amenity", "dentist")],
    "doctor": [("amenity", "doctors")],
    "lawyer": [("office", "lawyer")],
    "accountant": [("office", "accountant")],
    "real estate": [("office", "estate_agent")],
    "estate agent": [("office", "estate_agent")],
    "florist": [("shop", "florist")],
    "bookshop": [("shop", "books")],
    "electronics shop": [("shop", "electronics")],
    "clothing store": [("shop", "clothes")],
    "supermarket": [("shop", "supermarket")],
    "butcher": [("shop", "butcher")],
    "pharmacy": [("amenity", "pharmacy")],
    "pet shop": [("shop", "pet")],
    "furniture store": [("shop", "furniture")],
    "photographer": [("shop", "photo")],
    "driving school": [("amenity", "driving_school")],
    "tattoo studio": [("shop", "tattoo")],
    "veterinary": [("amenity", "veterinary")],
    "makeup studio": [("shop", "beauty")],
    "makeup artist": [("shop", "beauty")],
    "nail salon": [("shop", "beauty")],
    "spa": [("leisure", "spa")],
    "massage": [("shop", "massage")],
    "fashion designing": [("craft", "dressmaker"), ("shop", "tailor")],
    "fashion designer": [("craft", "dressmaker"), ("shop", "tailor")],
    "tailor": [("craft", "tailor"), ("shop", "tailor")],
    "dressmaker": [("craft", "dressmaker")],
    "leather production": [("craft", "leather")],
    "leatherworker": [("craft", "leather")],
    "shoemaker": [("craft", "shoemaker")],
    "cobbler": [("craft", "shoemaker")],
    "jewelry store": [("shop", "jewelry")],
    "jeweler": [("craft", "jeweler"), ("shop", "jewelry")],
    "bridal shop": [("shop", "bridal")],
    "boutique": [("shop", "boutique")],
    "shoe shop": [("shop", "shoes")],
    "carpenter": [("craft", "carpenter")],
    "plumber": [("craft", "plumber")],
    "electrician": [("craft", "electrician")],
    "painter": [("craft", "painter")],
    "welder": [("craft", "metal_construction")],
    "blacksmith": [("craft", "blacksmith")],
    "upholsterer": [("craft", "upholsterer")],
    "locksmith": [("craft", "locksmith"), ("shop", "locksmith")],
    "printer": [("craft", "printer")],
    "sign maker": [("craft", "signmaker")],
    "tiler": [("craft", "tiler")],
    "roofer": [("craft", "roofer")],
    "computer repair": [("shop", "computer")],
    "phone repair": [("shop", "mobile_phone")],
    "watch repair": [("craft", "watchmaker"), ("shop", "watches")],
    "physiotherapy": [("healthcare", "physiotherapist")],
    "optician": [("shop", "optician")],
    "fast food": [("amenity", "fast_food")],
    "ice cream shop": [("amenity", "ice_cream")],
    "confectionery": [("shop", "confectionery")],
    "daycare": [("amenity", "childcare")],
    "tutoring center": [("office", "educational_institution")],
    "music school": [("amenity", "music_school")],
    "event venue": [("amenity", "events_venue")],
    "catering": [("office", "catering")],
    "laundry": [("shop", "laundry")],
    "dry cleaner": [("shop", "dry_cleaning")],
    "travel agency": [("office", "travel_agent")],
    "insurance agency": [("office", "insurance")],
    "construction": [("craft", "builder"), ("office", "construction_company")],
    "interior design": [("office", "interior_design")],
    "cleaning company": [("office", "cleaning")],
    "security company": [("office", "security")],
    "photography studio": [("shop", "photo")],
    "event planning": [("office", "event_management")],
    "car dealership": [("shop", "car")],
}

# ---------------------------------------------------------------------------
# Regions -> countries
# ---------------------------------------------------------------------------
REGIONS = {
    "Africa": [
        "Nigeria", "Ghana", "Kenya", "South Africa", "Egypt", "Morocco",
        "Ethiopia", "Tanzania", "Uganda", "Senegal", "Ivory Coast",
        "Rwanda", "Zambia", "Cameroon", "Algeria",
    ],
    "Europe": [
        "United Kingdom", "France", "Germany", "Spain", "Italy", "Portugal",
        "Netherlands", "Belgium", "Switzerland", "Austria", "Ireland",
        "Poland", "Sweden", "Norway", "Denmark", "Finland", "Greece",
        "Czech Republic", "Romania", "Hungary", "Croatia", "Bulgaria",
    ],
    "North America": ["United States", "Canada", "Mexico"],
    "South America": ["Brazil", "Argentina", "Chile", "Colombia", "Peru", "Uruguay"],
    "Asia": [
        "India", "China", "Japan", "South Korea", "Indonesia", "Thailand",
        "Vietnam", "Philippines", "Malaysia", "Singapore", "Pakistan",
        "Bangladesh",
    ],
    "Middle East": [
        "United Arab Emirates", "Saudi Arabia", "Qatar", "Turkey", "Israel",
        "Jordan", "Kuwait",
    ],
    "Oceania": ["Australia", "New Zealand"],
}

COUNTRY_CALLING_CODES = {
    "United Kingdom": "44", "France": "33", "Germany": "49", "Spain": "34",
    "Italy": "39", "Portugal": "351", "Netherlands": "31", "Belgium": "32",
    "Switzerland": "41", "Austria": "43", "Ireland": "353", "Poland": "48",
    "Sweden": "46", "Norway": "47", "Denmark": "45", "Finland": "358",
    "Greece": "30", "Czech Republic": "420", "Romania": "40", "Hungary": "36",
    "Croatia": "385", "Bulgaria": "359",
    "Nigeria": "234", "Ghana": "233", "Kenya": "254", "South Africa": "27",
    "Egypt": "20", "Morocco": "212", "Ethiopia": "251", "Tanzania": "255",
    "Uganda": "256", "Senegal": "221", "Ivory Coast": "225", "Rwanda": "250",
    "Zambia": "260", "Cameroon": "237", "Algeria": "213",
    "United States": "1", "Canada": "1", "Mexico": "52",
    "Brazil": "55", "Argentina": "54", "Chile": "56", "Colombia": "57",
    "Peru": "51", "Uruguay": "598",
    "India": "91", "China": "86", "Japan": "81", "South Korea": "82",
    "Indonesia": "62", "Thailand": "66", "Vietnam": "84", "Philippines": "63",
    "Malaysia": "60", "Singapore": "65", "Pakistan": "92", "Bangladesh": "880",
    "United Arab Emirates": "971", "Saudi Arabia": "966", "Qatar": "974",
    "Turkey": "90", "Israel": "972", "Jordan": "962", "Kuwait": "965",
    "Australia": "61", "New Zealand": "64",
}

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def get_output_dir():
    termux_downloads = os.path.expanduser("~/storage/downloads")
    if os.path.isdir(termux_downloads):
        return termux_downloads
    print(
        "\n[Note] Can't see your Downloads folder. Run 'termux-setup-storage' once, "
        "then restart the script. Saving to current folder for now.\n"
    )
    return os.getcwd()


def normalize(text):
    return unicodedata.normalize("NFKD", text.strip().lower())


def resolve_niche(user_input):
    key = normalize(user_input)
    if key in NICHE_MAP:
        return key, NICHE_MAP[key]
    for name, tags in NICHE_MAP.items():
        if key in name or name in key:
            return name, tags
    return user_input.strip(), None


def parse_niches(raw_input):
    return [resolve_niche(c.strip()) for c in raw_input.split(",") if c.strip()]


def pick_from_menu(title, options, allow_multiple=False):
    print(f"\n{title}")
    for i, opt in enumerate(options, 1):
        print(f"  {i}. {opt}")
    prompt = "Enter number(s) comma separated (or 'all'): " if allow_multiple else "Enter number: "
    while True:
        raw = input(prompt).strip()
        if allow_multiple and raw.lower() == "all":
            return list(options)
        try:
            if allow_multiple:
                indices = [int(x.strip()) for x in raw.split(",") if x.strip()]
                if indices and all(1 <= i <= len(options) for i in indices):
                    return [options[i - 1] for i in indices]
            else:
                i = int(raw)
                if 1 <= i <= len(options):
                    return options[i - 1]
        except ValueError:
            pass
        print("Invalid — try again.")


def pick_regions_and_countries():
    region_names = list(REGIONS.keys())
    regions = pick_from_menu("Select continent(s):", region_names, allow_multiple=True)
    countries = []
    for region in regions:
        selected = pick_from_menu(f"Select country/countries in {region}:", REGIONS[region], allow_multiple=True)
        countries.extend(selected)
    seen, ordered = set(), []
    for c in countries:
        if c not in seen:
            seen.add(c)
            ordered.append(c)
    return ordered


# ---------------------------------------------------------------------------
# Geocoding
# ---------------------------------------------------------------------------

def geocode_place(place_name):
    params = {"q": place_name, "format": "json", "limit": 1}
    resp = SESSION.get(NOMINATIM_URL, params=params, timeout=30)
    resp.raise_for_status()
    results = resp.json()
    if not results:
        return None
    r = results[0]
    return {
        "osm_type": r.get("osm_type"),
        "osm_id": r.get("osm_id"),
        "bbox": [float(x) for x in r["boundingbox"]],
    }


# ---------------------------------------------------------------------------
# Overpass
# ---------------------------------------------------------------------------
GENERIC_CATEGORY_KEYS = "shop|amenity|office|craft|leisure|tourism|healthcare|sport"


def _keyword_to_regex(keyword):
    words = re.findall(r"[a-zA-Z0-9]+", keyword.lower())
    return "[ _-]*".join(re.escape(w) for w in words) if words else re.escape(keyword.lower())


def build_filter_clauses(tag_pairs, niche_label):
    if tag_pairs:
        return [f'nwr["{k}"="{v}"]' for k, v in tag_pairs]
    return [f'nwr[~"^({GENERIC_CATEGORY_KEYS})$"~"{_keyword_to_regex(niche_label)}",i]']


def build_overpass_query(tag_pairs, niche_label, geo, result_cap=500):
    clauses = build_filter_clauses(tag_pairs, niche_label)
    if geo["osm_type"] == "relation":
        area_id = 3600000000 + geo["osm_id"]
        parts = [f'{c}(area.searchArea);' for c in clauses]
        return f"[out:json][timeout:90];area({area_id})->.searchArea;({''.join(parts)});out center tags {result_cap};"
    else:
        s, n, w, e = geo["bbox"]
        parts = [f'{c}({s},{w},{n},{e});' for c in clauses]
        return f"[out:json][timeout:90];({''.join(parts)});out center tags {result_cap};"


def query_overpass(query, retries=3):
    for url in OVERPASS_URLS:
        for attempt in range(retries):
            try:
                resp = SESSION.post(url, data={"data": query}, timeout=100)
                if resp.status_code == 200:
                    return resp.json().get("elements", [])
                elif resp.status_code in (429, 504):
                    wait = 10 * (attempt + 1)
                    print(f"    Overpass busy ({resp.status_code}), retrying in {wait}s…")
                    time.sleep(wait)
                else:
                    break  # try next mirror
            except requests.RequestException as e:
                print(f"    Network error on {url}: {e}")
                time.sleep(3)
    return []


def has_website(tags):
    return any(tags.get(k) for k in ("website", "contact:website", "url"))


def get_phone(tags):
    for k in ("phone", "contact:phone", "contact:mobile"):
        if tags.get(k):
            return tags[k]
    return ""


def get_email(tags):
    for k in ("email", "contact:email"):
        if tags.get(k):
            return tags[k]
    return ""


def get_telegram(tags):
    for k in ("contact:telegram", "telegram"):
        val = tags.get(k)
        if val:
            return val if val.startswith("http") else f"https://t.me/{val.lstrip('@')}"
    return ""


def get_address(tags):
    parts = []
    street = tags.get("addr:street", "")
    housenumber = tags.get("addr:housenumber", "")
    if street:
        parts.append(f"{street} {housenumber}".strip())
    if tags.get("addr:city"):
        parts.append(tags["addr:city"])
    if tags.get("addr:postcode"):
        parts.append(tags["addr:postcode"])
    return ", ".join(parts) if parts else "N/A"


def make_wa_link(phone, country=None):
    if not phone:
        return ""
    phone = re.split(r"[;,]", phone)[0].strip()
    if not phone:
        return ""
    if phone.startswith("+"):
        digits = re.sub(r"\D", "", phone)
    elif phone.startswith("00"):
        digits = re.sub(r"\D", "", phone[2:])
    else:
        local = re.sub(r"\D", "", phone)
        code = COUNTRY_CALLING_CODES.get(country)
        if code:
            if local.startswith("0"):
                local = local[1:]
            digits = code + local
        else:
            digits = local
    return f"https://wa.me/{digits}" if digits else ""


def osm_leads_for_location(niche_label, tag_pairs, location_label, geo, country):
    query = build_overpass_query(tag_pairs, niche_label, geo)
    elements = query_overpass(query)
    print(f"    OSM found {len(elements)} total '{niche_label}' entries in {location_label}.")
    leads = []
    for el in elements:
        tags = el.get("tags", {})
        name = tags.get("name")
        if not name or has_website(tags):
            continue
        phone = get_phone(tags)
        osm_id = f'{el.get("type", "node")}/{el.get("id")}'
        leads.append({
            "name": name,
            "niche": niche_label,
            "country": country,
            "city": location_label if location_label != country else None,
            "address": get_address(tags),
            "phone": phone,
            "email": get_email(tags),
            "website": None,
            "facebook": tags.get("contact:facebook"),
            "instagram": tags.get("contact:instagram"),
            "telegram": get_telegram(tags),
            "wa_link": make_wa_link(phone, country),
            "google_rating": None,
            "review_count": None,
            "source": "osm",
            "source_id": osm_id,
        })
    return leads


# ---------------------------------------------------------------------------
# Google Places fallback
# ---------------------------------------------------------------------------

def google_leads_for_location(niche_label, location_label, country):
    if not GOOGLE_API_KEY:
        return []
    query = f"{niche_label} in {location_label}"
    params = {"query": query, "key": GOOGLE_API_KEY}
    results = []
    while len(results) < GOOGLE_MAX_RESULTS_PER_LOCATION:
        resp = SESSION.get(GOOGLE_TEXTSEARCH_URL, params=params, timeout=20)
        data = resp.json()
        if data.get("status") not in ("OK", "ZERO_RESULTS"):
            print(f"    Google error: {data.get('status')}")
            break
        results.extend(data.get("results", []))
        token = data.get("next_page_token")
        if not token or len(results) >= GOOGLE_MAX_RESULTS_PER_LOCATION:
            break
        params = {"pagetoken": token, "key": GOOGLE_API_KEY}
        time.sleep(2)

    results = results[:GOOGLE_MAX_RESULTS_PER_LOCATION]

    def fetch_details(place):
        try:
            r = SESSION.get(
                GOOGLE_DETAILS_URL,
                params={
                    "place_id": place["place_id"],
                    "fields": "website,formatted_phone_number,international_phone_number",
                    "key": GOOGLE_API_KEY,
                },
                timeout=15,
            )
            return r.json().get("result", {})
        except Exception:
            return {}

    with ThreadPoolExecutor(max_workers=GOOGLE_DETAIL_WORKERS) as pool:
        futures = {pool.submit(fetch_details, p): p for p in results}
        details_map = {id(f._args[0] if hasattr(f, "_args") else results[i]): future.result()
                       for i, future in enumerate(as_completed(futures))}

    leads = []
    for place, details in zip(results, [fetch_details(p) for p in results]):
        if details.get("website"):
            continue  # already has a website
        phone = details.get("international_phone_number") or details.get("formatted_phone_number", "")
        leads.append({
            "name": place.get("name"),
            "niche": niche_label,
            "country": country,
            "city": location_label if location_label != country else None,
            "address": place.get("formatted_address", "N/A"),
            "phone": phone,
            "email": None,
            "website": None,
            "facebook": None,
            "instagram": None,
            "telegram": None,
            "wa_link": make_wa_link(phone, country),
            "google_rating": place.get("rating"),
            "review_count": place.get("user_ratings_total"),
            "source": "google",
            "source_id": place.get("place_id"),
        })
    return leads


# ---------------------------------------------------------------------------
# LeadForge push
# ---------------------------------------------------------------------------

def push_to_leadforge(leads, niche_label):
    """POST leads to the LeadForge AI backend import endpoint."""
    if not LEADFORGE_URL or not LEADFORGE_TOKEN:
        print("\n[LeadForge] LEADFORGE_URL or LEADFORGE_TOKEN not set in CONFIG — skipping push.")
        return

    payload = {
        "businesses": [
            {
                "name": l["name"],
                "niche": l["niche"],
                "country": l["country"],
                "city": l.get("city"),
                "address": l.get("address"),
                "phone": l.get("phone"),
                "email": l.get("email"),
                "website": l.get("website"),
                "facebook": l.get("facebook"),
                "instagram": l.get("instagram"),
                "google_rating": l.get("google_rating"),
                "review_count": l.get("review_count"),
                "source_id": l.get("source_id"),
                "source": "termux_import",
            }
            for l in leads
        ]
    }

    print(f"\n[LeadForge] Pushing {len(leads)} leads to {LEADFORGE_URL}…")
    try:
        resp = SESSION.post(
            f"{LEADFORGE_URL}/api/leads/import",
            json=payload,
            headers={"X-Import-Token": LEADFORGE_TOKEN},
            timeout=30,
        )
        if resp.status_code == 200:
            r = resp.json()
            print(
                f"[LeadForge] ✓ Done! Added: {r['added']} | "
                f"Already in DB: {r['skipped_already_in_db']} | "
                f"Already emailed: {r['skipped_already_emailed']} | "
                f"Disqualified: {r['skipped_disqualified']}"
            )
            if r["added_names"]:
                print(f"[LeadForge] New leads: {', '.join(r['added_names'][:10])}"
                      + (" …and more" if len(r["added_names"]) > 10 else ""))
        elif resp.status_code == 401:
            print("[LeadForge] ✗ Auth failed — check your LEADFORGE_TOKEN in the CONFIG section.")
        elif resp.status_code == 503:
            print("[LeadForge] ✗ Import endpoint not configured on server — add IMPORT_API_KEY to Render.")
        else:
            print(f"[LeadForge] ✗ Error {resp.status_code}: {resp.text[:200]}")
    except requests.RequestException as e:
        print(f"[LeadForge] ✗ Network error: {e}")


# ---------------------------------------------------------------------------
# Word report
# ---------------------------------------------------------------------------

def write_report(all_leads, niche_label, output_dir):
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"LeadForge: {niche_label.title()} Leads")
    run.bold = True
    run.font.size = Pt(16)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(f"Generated {datetime.now().strftime('%Y-%m-%d %H:%M')} · {len(all_leads)} businesses")

    doc.add_paragraph()

    # Sort by review count descending, then rating
    sorted_leads = sorted(
        all_leads,
        key=lambda l: (l.get("review_count") or 0, l.get("google_rating") or 0),
        reverse=True,
    )

    for i, lead in enumerate(sorted_leads, 1):
        heading = doc.add_paragraph()
        heading.add_run(f"{i}. {lead['name']}").bold = True

        def row(label, value):
            if value and value != "N/A":
                p = doc.add_paragraph(style="List Bullet")
                r = p.add_run(f"{label}: ")
                r.bold = True
                p.add_run(str(value))

        row("Country", lead.get("country"))
        row("City", lead.get("city"))
        row("Address", lead.get("address"))
        row("Phone", lead.get("phone"))
        row("WhatsApp", lead.get("wa_link"))
        row("Email", lead.get("email"))
        row("Telegram", lead.get("telegram"))
        if lead.get("google_rating"):
            row("Rating", f"{lead['google_rating']} ⭐  ({lead.get('review_count', 0)} reviews)")
        row("Source", lead.get("source", "osm").upper())
        doc.add_paragraph()

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_niche = re.sub(r"[^a-zA-Z0-9_]", "_", niche_label)
    filename = os.path.join(output_dir, f"leads_{safe_niche}_{timestamp}.docx")
    doc.save(filename)
    return filename


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    print("\n=== LeadForge Lead Finder ===")
    print("Finds businesses WITHOUT a website, ready for your outreach.\n")

    raw_niches = input("Enter niche(s), comma separated (e.g. 'gym, bakery, dentist'): ")
    niches = parse_niches(raw_niches)
    if not niches:
        print("No niches entered. Exiting.")
        return

    countries = pick_regions_and_countries()

    # Optional city narrowing per country
    city_map = {}  # country -> list of cities (empty = whole country)
    for country in countries:
        narrow = input(f"\nNarrow {country} to specific cities? Enter city names comma-separated, or press Enter to search whole country: ").strip()
        city_map[country] = [c.strip() for c in narrow.split(",") if c.strip()]

    all_leads = []
    output_dir = get_output_dir()

    for niche_label, tag_pairs in niches:
        print(f"\n--- Searching: {niche_label} ---")
        for country in countries:
            cities = city_map.get(country) or [None]  # None = whole country
            for city in cities:
                location_label = city or country
                print(f"  Location: {location_label}")

                geo = geocode_place(location_label)
                if not geo:
                    print(f"    Could not geocode '{location_label}' — skipping.")
                    continue

                osm_results = osm_leads_for_location(
                    niche_label, tag_pairs, location_label, geo, country
                )
                all_leads.extend(osm_results)

                # Fall back to Google if OSM returned too few results
                if len(osm_results) < MIN_OSM_RESULTS_BEFORE_FALLBACK and GOOGLE_API_KEY:
                    print(f"    OSM thin — trying Google Places for {location_label}…")
                    google_results = google_leads_for_location(niche_label, location_label, country)
                    # Dedup against OSM results by name
                    existing_names = {l["name"].lower() for l in osm_results}
                    new_google = [g for g in google_results if g["name"].lower() not in existing_names]
                    print(f"    Google added {len(new_google)} extra leads.")
                    all_leads.extend(new_google)

    if not all_leads:
        print("\nNo leads found. Try different niches, locations, or add a Google API key.")
        return

    # Deduplicate by (name, country)
    seen = set()
    deduped = []
    for l in all_leads:
        key = (l["name"].lower(), l["country"])
        if key not in seen:
            seen.add(key)
            deduped.append(l)

    all_leads = sorted(deduped, key=lambda l: (l.get("review_count") or 0), reverse=True)
    niche_label_combined = ", ".join(n for n, _ in niches)

    print(f"\n✓ Found {len(all_leads)} unique leads with no website.")

    # Save Word report
    report_path = write_report(all_leads, niche_label_combined, output_dir)
    print(f"✓ Report saved: {report_path}")

    # Push to LeadForge
    if LEADFORGE_URL and LEADFORGE_TOKEN:
        if AUTO_PUSH:
            push_to_leadforge(all_leads, niche_label_combined)
        else:
            choice = input(f"\nPush {len(all_leads)} leads to LeadForge AI dashboard? (y/n): ").strip().lower()
            if choice == "y":
                push_to_leadforge(all_leads, niche_label_combined)
    else:
        print(
            "\n[Tip] Set LEADFORGE_URL and LEADFORGE_TOKEN in the CONFIG section at the top of this "
            "script to push leads directly to your LeadForge AI web app after each search."
        )

    print("\nDone!")


if __name__ == "__main__":
    main()
